
"""def new_func():
    f = open('t1.txt','w')
    f.write("hello world")
    print ("file created")
    l = ['a','e','i','o']
    print(l)
new_func()"""
import docx
import os
os.getcwd()
doc = docx.Document('report.docx')
doc.paragraphs
doc.paragraphs[0].text